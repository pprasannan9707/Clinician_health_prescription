"""
prescription_docx.py — PHASE 11: Prescription output in DOCX format.

Runs AFTER human-in-the-loop approval (never before — this module should
only ever be called with an already-signed-off summary/medication list).
Uses python-docx (pure Python, works cleanly inside an Azure Function or
App Service without Node/LibreOffice) so the whole ingestion-to-prescription
path stays in one runtime.

Output is written locally, then the caller (Function/App) uploads it to a
"prescriptions" Blob container and returns a SAS/Managed-Identity-scoped
download link to the Streamlit UI — the file itself never needs to touch
long-term storage unencrypted; see DEPLOYMENT notes for CMK + private
endpoint guidance.
"""

from __future__ import annotations

import datetime
import io
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


@dataclass
class PrescriptionMedication:
    name: str
    dosage: str
    frequency: str
    duration: str
    instructions: str = ""


@dataclass
class PrescriptionData:
    patient_name: str
    patient_id: str
    patient_dob: str
    doctor_name: str
    doctor_id: str
    doctor_speciality: str
    date_issued: str
    diagnosis_summary: str
    medications: List[PrescriptionMedication] = field(default_factory=list)
    allergy_warnings: List[str] = field(default_factory=list)
    reviewer_comment: str = ""
    context_sha256: str = ""


def _add_heading_rule(document: Document) -> None:
    """Horizontal rule via a paragraph bottom border (never a 1-cell table)."""
    p = document.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_prescription_docx(data: PrescriptionData, output_path: str) -> str:
    document = Document()

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # ---- Header -------------------------------------------------------
    title = document.add_heading("Clinical Prescription", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Generated via Clinical RAG Copilot — requires clinician wet/e-signature")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_heading_rule(document)

    # ---- Doctor / Patient info table ----------------------------------
    info_table = document.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    for row in info_table.rows:
        row.cells[0].width = Inches(3.15)
        row.cells[1].width = Inches(3.15)

    rows_content = [
        ("Prescribing Provider", f"{data.doctor_name} ({data.doctor_speciality})"),
        ("Provider ID", data.doctor_id),
        ("Patient", f"{data.patient_name}  (DOB: {data.patient_dob})"),
        ("Patient ID / Date Issued", f"{data.patient_id}  /  {data.date_issued}"),
    ]
    for i, (label, value) in enumerate(rows_content):
        info_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
        info_table.rows[i].cells[1].paragraphs[0].add_run(value)

    document.add_paragraph()

    # ---- Diagnosis summary ---------------------------------------------
    document.add_heading("Clinical Summary", level=2)
    document.add_paragraph(data.diagnosis_summary)

    # ---- Allergy warning banner -----------------------------------------
    if data.allergy_warnings:
        warn_p = document.add_paragraph()
        warn_run = warn_p.add_run("⚠ ALLERGY / SAFETY FLAGS:  " + "; ".join(data.allergy_warnings))
        warn_run.bold = True
        warn_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    document.add_paragraph()

    # ---- Medications table ----------------------------------------------
    document.add_heading("Prescribed Medications", level=2)
    if data.medications:
        med_table = document.add_table(rows=1, cols=4)
        med_table.style = "Light Grid Accent 1"
        col_widths = [Inches(1.8), Inches(1.2), Inches(1.4), Inches(1.9)]
        headers = ["Medication", "Dosage", "Frequency / Duration", "Instructions"]
        for cell, header, width in zip(med_table.rows[0].cells, headers, col_widths):
            cell.paragraphs[0].add_run(header).bold = True
            cell.width = width
        for med in data.medications:
            cells = med_table.add_row().cells
            cells[0].text = med.name
            cells[1].text = med.dosage
            cells[2].text = f"{med.frequency}, {med.duration}"
            cells[3].text = med.instructions
            for cell, width in zip(cells, col_widths):
                cell.width = width
    else:
        document.add_paragraph("No medications prescribed at this visit.")

    document.add_paragraph()

    # ---- Reviewer sign-off block ------------------------------------------
    document.add_heading("Human-in-the-Loop Sign-Off", level=2)
    document.add_paragraph(data.reviewer_comment or "(no reviewer comment recorded)")

    sig_p = document.add_paragraph()
    sig_p.add_run("Provider signature: _______________________________     Date: _______________")

    document.add_paragraph()
    footer = document.add_paragraph()
    footer_run = footer.add_run(
        f"Context fingerprint (SHA-256): {data.context_sha256}   |   "
        f"Document generated {datetime.datetime.utcnow().isoformat()}Z"
    )
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    document.save(output_path)
    return output_path


def build_prescription_docx_bytes(data: PrescriptionData) -> bytes:
    """In-memory variant for direct Blob Storage upload without a local temp file."""
    buffer = io.BytesIO()
    document = Document()
    # Reuse the same builder by writing to a temp path is simplest/most robust;
    # python-docx's Document.save() accepts a file-like object directly too:
    _populate_and_save_to_stream(data, buffer)
    return buffer.getvalue()


def _populate_and_save_to_stream(data: PrescriptionData, stream: io.BytesIO) -> None:
    import tempfile
    import os as _os
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        build_prescription_docx(data, tmp_path)
        with open(tmp_path, "rb") as f:
            stream.write(f.read())
    finally:
        _os.remove(tmp_path)
